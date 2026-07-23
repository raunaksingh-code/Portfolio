import os
import PyPDF2
from docx import Document
import pandas as pd

def extract_text():
    output = []
    
    # PDF
    try:
        reader = PyPDF2.PdfReader('Raunak Singh_CV.pdf')
        output.append(f"--- Raunak Singh_CV.pdf ---\n")
        output.append('\n'.join([page.extract_text() for page in reader.pages]))
    except Exception as e:
        output.append(f"Error reading CV: {e}")

    try:
        reader = PyPDF2.PdfReader('1774745664850.pdf')
        output.append(f"\n--- 1774745664850.pdf ---\n")
        output.append('\n'.join([page.extract_text() for page in reader.pages]))
    except Exception as e:
        pass

    try:
        reader = PyPDF2.PdfReader('Raunak Singh - P2511056 - SIP Report.pdf')
        output.append(f"\n--- SIP Report ---\n")
        # Just extract first few pages for context
        output.append('\n'.join([page.extract_text() for page in reader.pages[:10]]))
    except Exception as e:
        pass

    # DOCX
    try:
        doc = Document('BA Report.docx')
        output.append(f"\n--- BA Report.docx ---\n")
        output.append('\n'.join([p.text for p in doc.paragraphs][:50]))
    except Exception as e:
        pass

    try:
        doc = Document('Group 8.docx')
        output.append(f"\n--- Group 8.docx ---\n")
        output.append('\n'.join([p.text for p in doc.paragraphs][:50]))
    except Exception as e:
        pass

    try:
        doc = Document('Piramal Enterprises Ltd.docx')
        output.append(f"\n--- Piramal Enterprises Ltd.docx ---\n")
        output.append('\n'.join([p.text for p in doc.paragraphs][:50]))
    except Exception as e:
        pass

    with open('extracted_content.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(output))

if __name__ == '__main__':
    extract_text()
